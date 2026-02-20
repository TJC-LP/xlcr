#!/usr/bin/env python3
"""Generate a minimal DOCX for native-image tracing agent tests."""
import sys
from docx import Document

output = sys.argv[1] if len(sys.argv) > 1 else "test.docx"

doc = Document()
doc.add_heading("XLCR Test Document", level=1)
doc.add_paragraph(
    "This is a minimal test document used by the GraalVM tracing agent "
    "to capture Aspose.Words reflection metadata for native image builds."
)

# Add a table (exercises table rendering code paths)
table = doc.add_table(rows=3, cols=2)
table.style = "Table Grid"
table.cell(0, 0).text = "Format"
table.cell(0, 1).text = "Library"
table.cell(1, 0).text = "DOCX → PDF"
table.cell(1, 1).text = "Aspose.Words"
table.cell(2, 0).text = "DOC → PDF"
table.cell(2, 1).text = "Aspose.Words"

doc.save(output)
print(f"Generated {output}")
